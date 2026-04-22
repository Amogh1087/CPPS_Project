"""
AI Evaluation Module - Production-Ready Implementation
Uses Google Gemini API for semantic resume evaluation and python-docx for report generation.
"""

import os
import json
import logging
from typing import Dict, Optional
import google.generativeai as genai
from docx import Document

# 1. Load the hidden variables from the .env file
load_dotenv() 

# 2. Setup Professional Logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# 3. Securely grab the API key
API_KEY = os.getenv("GEMINI_API_KEY")

if not API_KEY:
    logger.error("GEMINI_API_KEY not found! Make sure your .env file is set up.")
else:
    genai.configure(api_key=API_KEY)

def generate_evaluation_report(parsed_resume_data: Dict, job_description: str, output_filename: str = "AI_Recommendations_Report.docx") -> Optional[str]:
    """
    Evaluates parsed resume data against a JD using LLM and generates a Word report.
    
    Args:
        parsed_resume_data (Dict): The structured dictionary returned by parser.py
        job_description (str): The target job description text
        output_filename (str): Desired name for the output .docx file
        
    Returns:
        Optional[str]: The file path to the saved document, or None if failed.
    """
    logger.info("Initializing Gemini API call...")
    
    # We use a system prompt that enforces strict JSON output
    model = genai.GenerativeModel('gemini-1.5-flash')
    
    # Convert the Python dictionary into a formatted JSON string for the prompt
    resume_string = json.dumps(parsed_resume_data, indent=2)
    
    prompt = f"""
    Act as an expert technical recruiter. Analyze this structured Resume Data against the Job Description.
    Provide the output STRICTLY as a valid JSON object with NO markdown formatting, NO backticks, and exactly these keys:
    "score_out_of_10": <float>,
    "interview_probability_percent": <integer>,
    "feedback": "<2-3 sentence technical summary>",
    "improved_bullet_points": ["<point 1>", "<point 2>", "<point 3>"]

    Resume Data:
    {resume_string}

    Job Description:
    {job_description}
    """

    try:
        # Call the LLM
        response = model.generate_content(prompt)
        
        # Clean the response to ensure it's pure JSON
        clean_text = response.text.strip().replace('```json', '').replace('```', '').strip()
        analysis_data = json.loads(clean_text)
        
        logger.info("LLM Evaluation successful. Compiling Word Document...")

        # Generate the Word Document
        doc = Document()
        doc.add_heading('AI Resume Analysis & Recommendation Report', 0)

        # Add candidate details from the parsed data
        candidate_name = parsed_resume_data.get("name", "Candidate")
        doc.add_paragraph(f"Prepared for: {candidate_name}")
        doc.add_paragraph("_" * 50) # visual divider

        # Add Metrics
        doc.add_heading('1. Suitability Metrics', level=1)
        doc.add_paragraph(f"Overall Match Score: {analysis_data.get('score_out_of_10', 'N/A')}/10")
        doc.add_paragraph(f"Interview Probability: {analysis_data.get('interview_probability_percent', 'N/A')}%")

        # Add Feedback
        doc.add_heading('2. Recruiter Feedback', level=1)
        doc.add_paragraph(analysis_data.get('feedback', 'No feedback provided.'))

        # Add Bullet Points
        doc.add_heading('3. Recommended Resume Updates', level=1)
        doc.add_paragraph("Consider replacing weak points in your experience section with these optimized bullets:")
        for point in analysis_data.get('improved_bullet_points', []):
            doc.add_paragraph(point, style='List Bullet')

        doc.save(output_filename)
        logger.info(f"Report successfully saved to {output_filename}")
        
        return output_filename

    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse LLM response into JSON: {e}")
        logger.debug(f"Raw LLM response was: {response.text}")
        return None
    except Exception as e:
        logger.error(f"An unexpected error occurred during evaluation: {e}")
        return None